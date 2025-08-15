#!/usr/bin/env python
"""Test WhatsApp integration"""

from modules.whatsapp_utils import send_text_message, send_pdf_to_whatsapp
from modules.pdf_utils import generate_pdf
import perplexity_config

from docx import Document

def generate_prescription(data):
    doc = Document()
    doc.add_heading('Prescription', level=1)

    doc.add_paragraph(f"Patient Name: {data['patient_name']}")
    doc.add_paragraph(f"Age: {data['age']}")
    doc.add_paragraph(f"Date: {data['date']}")
    doc.add_paragraph(f"Diagnosis: {data['diagnosis']}")
    doc.add_paragraph("Medications:")
    for med in data['medications']:
        doc.add_paragraph(f"- {med}", style='List Bullet')

    doc.add_paragraph(f"Instructions: {data['instructions']}")
    doc.add_paragraph(f"Doctor: {data['doctor_name']} ({data['doctor_id']})")

    doc.save("prescription.docx")


def test_whatsapp():
    """Test WhatsApp functionality"""
    print("Testing WhatsApp integration...")
    
    # Get config
    access_token = perplexity_config.WHATSAPP_ACCESS_TOKEN
    phone_id = getattr(perplexity_config, 'META_WA_PHONE_NUMBER_ID', '783828511470086')
    
    print(f"Access Token: {access_token[:10]}...")
    print(f"Phone ID: {phone_id}")
    
    # Test phone number (use your own for testing)
    test_phone = "6363738550"  # Replace with your test number
    
    # Test text message
    print("\n1. Testing text message...")
    result = send_text_message(
        test_phone,
        "Test message from MauEyeCare system",
        access_token,
        phone_id
    )
    print(f"Text result: {result}")
    
    # Test PDF generation and sending
    print("\n2. Testing PDF generation...")
    pdf_bytes = generate_pdf(
        {"Eye Drops": 1}, "", "", "Dr Danish", 
        "Test Patient", 30, "Male", "Test advice", 
        {"OD": {"Sphere": "+1.00"}}, ["Single Vision"]
    )
    
    if isinstance(pdf_bytes, bytes) and len(pdf_bytes) > 100:
        print(f"PDF generated: {len(pdf_bytes)} bytes")
        
        print("\n3. Testing PDF send via WhatsApp...")
        result = send_pdf_to_whatsapp(
            test_phone,
            pdf_bytes,
            "Test Patient",
            access_token,
            phone_id
        )
        print(f"PDF send result: {result}")
    else:
        print("PDF generation failed")

if __name__ == "__main__":
    test_whatsapp()