"""
Simple DOCX Generator - Reliable prescription generation
"""
from io import BytesIO
import datetime

def generate_simple_prescription_docx(prescription, doctor_name, patient_name, age, gender, advice, rx_table, recommendations, dosages=None):
    """Generate simple but reliable DOCX prescription"""
    
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create document
        doc = Document()
        
        # Header
        header = doc.add_heading('MauEyeCare Optical Center', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Clinic info
        clinic_para = doc.add_paragraph()
        clinic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clinic_para.add_run('Dr. Danish - Eye Care Specialist\n').bold = True
        clinic_para.add_run('Phone: +91 92356-47410 | Email: info@maueyecare.com')
        
        # Date
        timestamp = datetime.datetime.now()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(f'Date: {timestamp.strftime("%d/%m/%Y")}\n')
        date_para.add_run(f'Prescription No: RX-{timestamp.strftime("%Y%m%d%H%M")}')
        
        # Patient info
        doc.add_paragraph(f'Patient Name: {patient_name}')
        doc.add_paragraph(f'Age: {age} years | Gender: {gender}')
        
        # RX Details
        if rx_table:
            doc.add_heading('Prescription Details:', level=2)
            for eye in ['OD', 'OS']:
                eye_data = rx_table.get(eye, {})
                if eye_data.get('Sphere'):
                    rx_line = f"{eye}: "
                    if eye_data.get('Sphere'):
                        rx_line += f"SPH {eye_data['Sphere']} "
                    if eye_data.get('Cylinder'):
                        rx_line += f"CYL {eye_data['Cylinder']} "
                    if eye_data.get('Axis'):
                        rx_line += f"AXIS {eye_data['Axis']}"
                    doc.add_paragraph(rx_line)
        
        # Medicines
        if prescription:
            doc.add_heading('Prescribed Medications:', level=2)
            for item, qty in prescription.items():
                med_para = doc.add_paragraph()
                med_para.add_run(f'• {item} - Quantity: {qty}').bold = True
                
                if dosages and item in dosages:
                    dosage_info = dosages[item]
                    med_para.add_run(f'\n  Dosage: {dosage_info.get("dosage", "As directed")}')
                    med_para.add_run(f'\n  Timing: {dosage_info.get("timing", "As directed")}')
        
        # Advice
        if advice:
            doc.add_heading('Doctor\'s Advice:', level=2)
            doc.add_paragraph(advice)
        
        # Instructions
        doc.add_heading('Instructions:', level=2)
        instructions = [
            'Take medications as prescribed',
            'Follow up if symptoms persist',
            'Avoid rubbing eyes',
            'Contact clinic for emergencies'
        ]
        for instruction in instructions:
            doc.add_paragraph(f'• {instruction}')
        
        # Signature
        signature_para = doc.add_paragraph()
        signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature_para.add_run('\n\nDr. Danish\nEye Care Specialist')
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        # Fallback to text format
        return generate_text_prescription(prescription, doctor_name, patient_name, age, gender, advice, rx_table, recommendations, dosages)

def generate_text_prescription(prescription, doctor_name, patient_name, age, gender, advice, rx_table, recommendations, dosages=None):
    """Generate text-based prescription as fallback"""
    
    timestamp = datetime.datetime.now()
    
    content = f"""MauEyeCare Optical Center
Dr. Danish - Eye Care Specialist
Phone: +91 92356-47410 | Email: info@maueyecare.com

Date: {timestamp.strftime("%d/%m/%Y")}
Prescription No: RX-{timestamp.strftime("%Y%m%d%H%M")}

Patient Name: {patient_name}
Age: {age} years | Gender: {gender}

"""
    
    # RX Details
    if rx_table:
        content += "Prescription Details:\n"
        for eye in ['OD', 'OS']:
            eye_data = rx_table.get(eye, {})
            if eye_data.get('Sphere'):
                rx_line = f"{eye}: "
                if eye_data.get('Sphere'):
                    rx_line += f"SPH {eye_data['Sphere']} "
                if eye_data.get('Cylinder'):
                    rx_line += f"CYL {eye_data['Cylinder']} "
                if eye_data.get('Axis'):
                    rx_line += f"AXIS {eye_data['Axis']}"
                content += rx_line + "\n"
        content += "\n"
    
    # Medicines
    if prescription:
        content += "Prescribed Medications:\n"
        for item, qty in prescription.items():
            content += f"• {item} - Quantity: {qty}\n"
            if dosages and item in dosages:
                dosage_info = dosages[item]
                content += f"  Dosage: {dosage_info.get('dosage', 'As directed')}\n"
                content += f"  Timing: {dosage_info.get('timing', 'As directed')}\n"
        content += "\n"
    
    # Advice
    if advice:
        content += f"Doctor's Advice:\n{advice}\n\n"
    
    # Instructions
    content += """Instructions:
• Take medications as prescribed
• Follow up if symptoms persist
• Avoid rubbing eyes
• Contact clinic for emergencies

                                Dr. Danish
                                Eye Care Specialist
"""
    
    return content.encode('utf-8')