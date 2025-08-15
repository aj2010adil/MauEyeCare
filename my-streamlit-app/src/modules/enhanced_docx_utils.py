#!/usr/bin/env python
"""
Enhanced DOCX prescription generation with professional features
"""
try:
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    print("Warning: python-docx not available. DOCX generation will be disabled.")

__all__ = ['generate_professional_prescription_docx', 'generate_inventory_report_docx', 'DOCX_AVAILABLE']

from io import BytesIO
import datetime

def generate_professional_prescription_docx(prescription, doctor_name, patient_name, age, gender, advice, rx_table, recommendations, dosages=None):
    """Generate professional prescription in DOCX format with enhanced features"""
    
    if not DOCX_AVAILABLE or Document is None:
        # Return a simple text-based prescription as fallback
        return generate_text_prescription_fallback(prescription, doctor_name, patient_name, age, gender, advice, rx_table, recommendations, dosages)
    
    doc = Document()
    
    # Header with logo
    header = doc.add_heading('👁️ MauEyeCare Optical Center', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Clinic details
    clinic_info = doc.add_paragraph()
    clinic_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clinic_info.add_run('Dr. Danish - Eye Care Specialist\n').bold = True
    clinic_info.add_run('Reg. No: MCI-12345 | Phone: +91 92356-47410\n')
    clinic_info.add_run('Email: info@maueyecare.com')
    
    # Separator
    doc.add_paragraph('═' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date and prescription number
    timestamp = datetime.datetime.now()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f'Date: {timestamp.strftime("%d/%m/%Y %H:%M")}\n')
    date_para.add_run(f'Prescription No: RX-{timestamp.strftime("%Y%m%d%H%M")}')
    
    # Patient info
    doc.add_paragraph(f'Patient Name: {patient_name}')
    doc.add_paragraph(f'Age: {age} years    Gender: {gender}')
    
    # RX Table
    if rx_table and any(rx_table.get(eye, {}).get('Sphere') for eye in ['OD', 'OS']):
        doc.add_heading('Prescription Details:', level=2)
        
        table = doc.add_table(rows=3, cols=8)
        table.style = 'Table Grid'
        
        headers = ['Eye', 'Sphere', 'Cylinder', 'Axis', 'Prism', 'Near Vision', 'Glass Type', 'Glass Tint']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        for row_idx, eye in enumerate(['OD', 'OS'], 1):
            eye_data = rx_table.get(eye, {})
            table.cell(row_idx, 0).text = eye
            table.cell(row_idx, 1).text = eye_data.get('Sphere', '')
            table.cell(row_idx, 2).text = eye_data.get('Cylinder', '')
            table.cell(row_idx, 3).text = eye_data.get('Axis', '')
            table.cell(row_idx, 4).text = eye_data.get('Prism', '')
            table.cell(row_idx, 5).text = eye_data.get('NearVision', '')
            table.cell(row_idx, 6).text = eye_data.get('GlassType', '')
            table.cell(row_idx, 7).text = eye_data.get('GlassTint', '')
    
    # Medicines with dosages
    if prescription:
        doc.add_heading('Prescribed Medications:', level=2)
        
        med_table = doc.add_table(rows=len(prescription)+1, cols=4)
        med_table.style = 'Table Grid'
        
        # Headers
        headers = ['Medicine', 'Quantity', 'Dosage', 'Timing']
        for i, header in enumerate(headers):
            cell = med_table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        # Medicine data
        for idx, (item, qty) in enumerate(prescription.items(), 1):
            med_table.cell(idx, 0).text = item
            med_table.cell(idx, 1).text = str(qty)
            
            # Get dosage info
            if dosages and item in dosages:
                dosage_info = dosages[item]
                med_table.cell(idx, 2).text = dosage_info.get('dosage', '1 drop')
                med_table.cell(idx, 3).text = dosage_info.get('timing', '2 times daily')
            else:
                med_table.cell(idx, 2).text = '1 drop' if 'drop' in item.lower() else '1 tablet'
                med_table.cell(idx, 3).text = '2 times daily'
    
    # Recommendations
    if recommendations:
        doc.add_heading('Recommendations:', level=2)
        for rec in recommendations:
            doc.add_paragraph(f'• {rec}', style='List Bullet')
    
    # Advice
    if advice:
        doc.add_heading('Doctor\'s Advice:', level=2)
        doc.add_paragraph(advice)
    
    # Instructions
    doc.add_heading('General Instructions:', level=2)
    instructions = [
        'Take medications as prescribed',
        'Follow up after 2 weeks if symptoms persist',
        'Avoid rubbing eyes',
        'Use prescribed eye drops regularly',
        'Contact clinic for any emergency'
    ]
    for instruction in instructions:
        doc.add_paragraph(f'• {instruction}', style='List Bullet')
    
    # Footer
    doc.add_paragraph('═' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = doc.add_paragraph('Thank you for choosing MauEyeCare Optical Center')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Doctor signature
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.add_run('\n\nDr. Danish\nEye Care Specialist')
    
    # Convert to bytes
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer.getvalue()

def generate_inventory_report_docx(inventory_data):
    """Generate professional inventory report"""
    
    if not DOCX_AVAILABLE or Document is None:
        return generate_text_inventory_fallback(inventory_data)
    
    doc = Document()
    
    # Header
    header = doc.add_heading('📦 MauEyeCare Inventory Report', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    timestamp = datetime.datetime.now()
    date_para = doc.add_paragraph(f'Generated on: {timestamp.strftime("%d/%m/%Y %H:%M")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Inventory table
    if inventory_data:
        table = doc.add_table(rows=len(inventory_data)+1, cols=4)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Item Name', 'Current Stock', 'Status', 'Action Required']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        # Data
        for idx, (item, qty) in enumerate(inventory_data.items(), 1):
            table.cell(idx, 0).text = item
            table.cell(idx, 1).text = str(qty)
            
            if qty == 0:
                table.cell(idx, 2).text = 'Out of Stock'
                table.cell(idx, 3).text = 'Reorder Immediately'
            elif qty < 5:
                table.cell(idx, 2).text = 'Low Stock'
                table.cell(idx, 3).text = 'Reorder Soon'
            else:
                table.cell(idx, 2).text = 'In Stock'
                table.cell(idx, 3).text = 'No Action'
    
    # Summary
    doc.add_heading('Summary:', level=2)
    total_items = len(inventory_data)
    low_stock = sum(1 for qty in inventory_data.values() if qty < 5)
    out_of_stock = sum(1 for qty in inventory_data.values() if qty == 0)
    
    doc.add_paragraph(f'Total Items: {total_items}')
    doc.add_paragraph(f'Low Stock Items: {low_stock}')
    doc.add_paragraph(f'Out of Stock Items: {out_of_stock}')
    
    # Convert to bytes
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer.getvalue()

def generate_text_prescription_fallback(prescription, doctor_name, patient_name, age, gender, advice, rx_table, recommendations, dosages=None):
    """Fallback text-based prescription when DOCX is not available"""
    
    timestamp = datetime.datetime.now()
    
    text_content = f"""
👁️ MauEyeCare Optical Center
Dr. Danish - Eye Care Specialist
Reg. No: UPS 2908 | Phone: +91 92356-47410
Email: info@maueyecare.com

{'='*60}

Date: {timestamp.strftime("%d/%m/%Y %H:%M")}
Prescription No: RX-{timestamp.strftime("%Y%m%d%H%M")}

Patient Name: {patient_name}
Age: {age} years    Gender: {gender}

"""
    
    # RX Table
    if rx_table and any(rx_table.get(eye, {}).get('Sphere') for eye in ['OD', 'OS']):
        text_content += "\nPrescription Details:\n"
        text_content += "-" * 40 + "\n"
        for eye in ['OD', 'OS']:
            eye_data = rx_table.get(eye, {})
            if eye_data.get('Sphere'):
                text_content += f"{eye}: SPH {eye_data.get('Sphere', '')} CYL {eye_data.get('Cylinder', '')} AXIS {eye_data.get('Axis', '')}\n"
    
    # Medicines
    if prescription:
        text_content += "\nPrescribed Medications:\n"
        text_content += "-" * 40 + "\n"
        for item, qty in prescription.items():
            dosage_info = dosages.get(item, {}) if dosages else {}
            dosage = dosage_info.get('dosage', '1 drop' if 'drop' in item.lower() else '1 tablet')
            timing = dosage_info.get('timing', '2 times daily')
            text_content += f"• {item} - Qty: {qty} - {dosage} {timing}\n"
    
    # Recommendations
    if recommendations:
        text_content += "\nRecommendations:\n"
        text_content += "-" * 40 + "\n"
        for rec in recommendations:
            text_content += f"• {rec}\n"
    
    # Advice
    if advice:
        text_content += f"\nDoctor's Advice:\n{advice}\n"
    
    text_content += f"""

General Instructions:
• Take medications as prescribed
• Follow up after 2 weeks if symptoms persist
• Avoid rubbing eyes
• Use prescribed eye drops regularly
• Contact clinic for any emergency

{'='*60}
Thank you for choosing MauEyeCare Optical Center

                                    Dr. Danish
                                    Eye Care Specialist
"""
    
    return text_content.encode('utf-8')

def generate_text_inventory_fallback(inventory_data):
    """Fallback text-based inventory report"""
    
    timestamp = datetime.datetime.now()
    
    text_content = f"""
📦 MauEyeCare Inventory Report
Generated on: {timestamp.strftime("%d/%m/%Y %H:%M")}

{'='*60}

Inventory Status:
{'-'*40}
"""
    
    if inventory_data:
        for item, qty in inventory_data.items():
            if qty == 0:
                status = "Out of Stock - Reorder Immediately"
            elif qty < 5:
                status = "Low Stock - Reorder Soon"
            else:
                status = "In Stock"
            
            text_content += f"{item}: {qty} units - {status}\n"
    
    # Summary
    total_items = len(inventory_data)
    low_stock = sum(1 for qty in inventory_data.values() if qty < 5)
    out_of_stock = sum(1 for qty in inventory_data.values() if qty == 0)
    
    text_content += f"""

Summary:
Total Items: {total_items}
Low Stock Items: {low_stock}
Out of Stock Items: {out_of_stock}

{'='*60}
"""
    
    return text_content.encode('utf-8')