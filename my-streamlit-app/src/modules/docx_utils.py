#!/usr/bin/env python
"""
DOCX prescription generation utilities
"""
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
import datetime

def generate_prescription_docx(prescription, doctor_name, patient_name, age, gender, advice, rx_table, recommendations):
    """Generate professional prescription in DOCX format"""
    
    doc = Document()
    
    # Header
    header = doc.add_heading('MauEyeCare Optical Center', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Doctor info
    doc.add_paragraph('Dr. Danish - Eye Care Specialist', style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Phone: +91 92356-47410', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('_' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date and patient info
    date_para = doc.add_paragraph(f'Date: {datetime.datetime.now().strftime("%d/%m/%Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph(f'Patient Name: {patient_name}')
    doc.add_paragraph(f'Age: {age} years    Gender: {gender}')
    
    # RX Table if available
    if rx_table and any(rx_table.get(eye, {}).get('Sphere') for eye in ['OD', 'OS']):
        doc.add_heading('Prescription Details:', level=2)
        
        table = doc.add_table(rows=3, cols=8)
        table.style = 'Table Grid'
        
        # Headers
        headers = ['Eye', 'Sphere', 'Cylinder', 'Axis', 'Prism', 'Near Vision', 'Glass Type', 'Glass Tint']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        # OD and OS data
        for row_idx, eye in enumerate(['OD', 'OS'], 1):
            eye_data = rx_table.get(eye, {})
            table.cell(row_idx, 0).text = eye
            table.cell(row_idx, 1).text = eye_data.get('Sphere', '')
            table.cell(row_idx, 2).text = eye_data.get('Cylinder', '')
            table.cell(row_idx, 3).text = eye_data.get('Axis', '')
            table.cell(row_idx, 4).text = eye_data.get('Prism', '')
            table.cell(row_idx, 5).text = eye_data.get('NearVision', '')
            table.cell(row_idx, 6).text = eye_data.get('GlassType', '')
            table.cell(row_idx, 7).text = eye_data.get('GlassTint', '')
    
    # Medicines/Items
    if prescription:
        doc.add_heading('Prescribed Items:', level=2)
        for item, qty in prescription.items():
            doc.add_paragraph(f'• {item} - Quantity: {qty}', style='List Bullet')
    
    # Recommendations
    if recommendations:
        doc.add_heading('Recommendations:', level=2)
        for rec in recommendations:
            doc.add_paragraph(f'• {rec}', style='List Bullet')
    
    # Advice
    if advice:
        doc.add_heading('Doctor\'s Advice:', level=2)
        doc.add_paragraph(advice)
    
    # Footer
    doc.add_paragraph('_' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = doc.add_paragraph('Thank you for choosing MauEyeCare Optical Center')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Convert to bytes
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer.getvalue()